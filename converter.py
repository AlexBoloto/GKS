# -*- coding: utf8 -*-

import csv
import xlwt
import openpyxl

from docx import Document
from docx.table import _Cell
from docx.oxml.simpletypes import ST_Merge


def __extract_table(table):
    """Extracts table data from table object"""
    results = []
    n = 0
    for tr in table._tbl.tr_lst:
        r = []
        for tc in tr.tc_lst:
            for grid_span_idx in range(tc.grid_span):
                if tc.vMerge == ST_Merge.CONTINUE:
                    r.append(results[n - 1][len(r) - 1])
                elif grid_span_idx > 0:
                    r.append(r[-1])
                else:
                    cell = _Cell(tc, table)
                    r.append(cell.text.replace('\n', ' ').encode('utf8'))
        results.append(r)
        n += 1
    return results

def __store_table(tabdata, filename, format='csv'):
    """Saves table data as csv file"""
    if format == 'csv':
        f = open(filename, 'w')
        w = csv.writer(f, delimiter=',')
        for row in tabdata:
            w.writerow(row)
    elif format == 'xls':
        workbook = xlwt.Workbook()
        ws = __xls_table_to_sheet(tabdata, workbook.add_sheet('0'))
        print(dir(ws))
        workbook.save(filename)
    elif format == 'xlsx':
        workbook = openpyxl.Workbook()
        ws = __xlsx_table_to_sheet(tabdata, workbook.create_sheet('0'))
        workbook.save(filename)


def __xls_table_to_sheet(table, ws):
    rn = 0
    for row in table:
        cn = 0
        for c in row:
            ws.write(rn, cn, c.decode('utf8'))
            cn += 1
        rn += 1
    return ws

def __xlsx_table_to_sheet(table, ws):
    rn = 0
    for row in table:
        ws.append(row)
        rn += 1
    return ws



def extract_tables(filename):
    """Extracts table from .DOCX files"""
    tables = []
    document = Document(filename)
    n = 0
    for table in document.tables:
        n += 1
        tdata = __extract_table(table)
        tables.append(tdata)
    return tables


def extract(filename, format='csv', sizefilter=0, singlefile=False):
    """Extracts tables from csv files and saves them as csv, xls or xlsx files"""
    tables = extract_tables(filename)
    name = filename.rsplit('.', 1)[0]
    format = format.lower()
    n = 0
    lfilter = int(sizefilter)
    if singlefile:
        if format == 'xls':
            workbook = xlwt.Workbook()
            for t in tables:
                if lfilter >= len(t):
                    continue
                n += 1
                ws = __xls_table_to_sheet(t, workbook.add_sheet(str(n)))
            destname = name + '.%s' % (format)
            workbook.save(destname)
        elif format == 'xlsx':
            workbook = openpyxl.Workbook()
            for t in tables:
                if lfilter >= len(t):
                    continue
                n += 1
                #ws = __xlsx_table_to_sheet(t, workbook.create_sheet(str(n)))
                ws = __xlsx_table_to_sheet(t, workbook.active)
            #destname = name + '.%s' % (format)
            ws = add_table_name(ws, str(name[len(name)-7:]))
            for i in range(5, ws.max_column):
                delete_column(ws, i)
            destname = name[:len(name)-21] + 'excel\\' + name[len(name)-7:] + '.%s' % (format)
            workbook.save(destname)
    else:
        for t in tables:
            if lfilter >= len(t):
                continue
            n += 1
            destname = name + '_%d.%s' % (n, format)
            __store_table(t, destname, format)

def add_table_name(ws, table_name):
    for i in range(1, ws.max_row + 1):
        cellref = ws.cell(row=i, column=4)
        cellref.number_format = 'General'
        cellref.value = str(table_name)
    return ws

def delete_column(ws, delete_column):
    if isinstance(delete_column, str):
        delete_column = openpyxl.utils.column_index_from_string(delete_column)
    assert delete_column >= 1, "Column numbers must be 1 or greater"
    for column in range(delete_column, ws.max_column + 1):
        for row in range(1, ws.max_row + 1):
            ws.cell(row=row, column=column).value = ''